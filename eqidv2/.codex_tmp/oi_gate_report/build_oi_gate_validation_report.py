from __future__ import annotations

import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
DEPS = ROOT / ".codex_tmp" / "oi_gate_report" / "deps"
sys.path.insert(0, str(DEPS))

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = ROOT / "reports" / "OI_GATE_VALIDATION_V6_V10_V12_20260902.docx"
QA = ROOT / ".codex_tmp" / "oi_gate_report" / "qa"
QA.mkdir(parents=True, exist_ok=True)
OUT.parent.mkdir(parents=True, exist_ok=True)

# standard_business_brief preset, with named overrides:
# - title 24 pt / dark navy for a technical memo masthead
# - table text 8.5-9 pt where numeric density requires it
# - compact figure captions at 8.5 pt
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF4D6"
PALE_RED = "FBEAEA"
GREEN = "2F6B3C"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
TABLE_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9DEE5", size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, *, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != TABLE_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, *, size=9, bold=False, color=INK, align=None):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if align is not None:
            p.alignment = align
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_dxa, *, numeric_cols=(), header_fill=LIGHT, size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        hdr.cells[idx].text = str(text)
        set_cell_shading(hdr.cells[idx], header_fill)
        style_cell_text(
            hdr.cells[idx],
            size=size,
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
        )
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        for idx, value in enumerate(values):
            row.cells[idx].text = str(value)
            if r_idx % 2 == 1:
                set_cell_shading(row.cells[idx], "FAFBFC")
            style_cell_text(
                row.cells[idx],
                size=size,
                align=WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_table_geometry(table, widths_dxa)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    return table


def add_callout(doc, label, text, *, fill=PALE_BLUE, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), "D7DEE8")
        borders.append(node)
    p_pr.append(borders)
    r = p.add_run(label + "  ")
    set_run_font(r, size=11, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_para(doc, text="", *, bold=False, italic=False, color=INK, size=11, align=None, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich_para(doc, parts, *, after=6, align=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = keep
    if align is not None:
        p.alignment = align
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run_font(r, **kwargs)
    return p


def add_bullet(doc, text, *, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(
        r,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
    )
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_run_font(r, size=8.5, color=MUTED, italic=True)
    if not p.runs:
        set_run_font(p.add_run(text), size=8.5, color=MUTED, italic=True)
    else:
        p.runs[0].text = text
    return p


def set_last_image_alt(doc, description):
    drawings = doc._element.xpath(".//wp:docPr")
    if drawings:
        drawings[-1].set("descr", description)
        drawings[-1].set("title", description[:120])


def add_picture(doc, path, caption, alt, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    set_last_image_alt(doc, alt)
    add_caption(doc, caption)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=8.5, color=MUTED)


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ):
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Tune real built-in list styles to the selected preset.
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167
    if "List Bullet 2" in styles:
        st = styles["List Bullet 2"]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.75)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.167

    cap = styles["Caption"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = rgb(MUTED)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("OI GATE VALIDATION  |  V6 -> V10 -> V12"), size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("Independent technical review  |  "), size=8.5, color=MUTED)
    add_page_field(fp)
    return doc


def charts():
    months = ["May", "Jun", "Jul", "Aug"]
    prev_oi = np.array([16_987.5, 135_750, 1_944_800, 11_583_000])
    oi_pct = np.array([7.275132, 1.694915, 0.882658, 0.430408])
    fig, ax1 = plt.subplots(figsize=(9.2, 4.5))
    bars = ax1.bar(months, prev_oi, color="#9FBAD0", width=0.58)
    ax1.set_yscale("log")
    ax1.set_ylabel("Median previous OI (log scale)", color="#17365D")
    ax1.grid(axis="y", alpha=0.2)
    for bar, val in zip(bars, prev_oi):
        ax1.text(bar.get_x() + bar.get_width()/2, val*1.15, f"{val:,.0f}", ha="center", va="bottom", fontsize=8, color="#17365D")
    ax2 = ax1.twinx()
    ax2.plot(months, oi_pct, color="#9B1C1C", marker="o", linewidth=2.4)
    ax2.set_ylabel("Median 5-minute OI change (%)", color="#9B1C1C")
    ax2.set_ylim(0, 8.4)
    for x, val in zip(months, oi_pct):
        ax2.text(x, val + 0.28, f"{val:.2f}%", ha="center", fontsize=8, color="#9B1C1C")
    ax1.set_title("V6 selected trades: OI base rises while percentage change falls", loc="left", fontweight="bold", color="#17365D")
    fig.tight_layout()
    path1 = QA / "monthly_oi_drift.png"
    fig.savefig(path1, dpi=180, bbox_inches="tight")
    plt.close(fig)

    labels = ["Frozen\nthresholds", "All leg thresholds\nset to 0.05%"]
    net = [98.574007, 75.297490]
    pf = [2.108901, 1.653741]
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.1))
    colors = ["#2E74B5", "#C99A2E"]
    axes[0].bar(labels, net, color=colors, width=0.6)
    axes[0].set_title("Net return points", fontweight="bold", color="#17365D")
    axes[0].set_ylim(0, 112)
    for i, v in enumerate(net): axes[0].text(i, v+2, f"{v:.2f}", ha="center", fontsize=9)
    axes[1].bar(labels, pf, color=colors, width=0.6)
    axes[1].set_title("Profit factor", fontweight="bold", color="#17365D")
    axes[1].axhline(1.0, color="#9B1C1C", linestyle="--", linewidth=1)
    axes[1].set_ylim(0, 2.45)
    for i, v in enumerate(pf): axes[1].text(i, v+0.05, f"{v:.3f}", ha="center", fontsize=9)
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.18)
    fig.suptitle("V6 per-leg OI threshold ablation (base OI gate still active)", x=0.06, ha="left", fontweight="bold", color="#17365D")
    fig.tight_layout(rect=[0, 0, 1, 0.92])
    path2 = QA / "v6_threshold_ablation.png"
    fig.savefig(path2, dpi=180, bbox_inches="tight")
    plt.close(fig)

    labels = ["Passed high OI\nthreshold", "Failed high OI\nthreshold"]
    net = [12.863805, -13.612536]
    pf = [1.788142, 0.429258]
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.1))
    colors = ["#3F7D52", "#B65454"]
    axes[0].bar(labels, net, color=colors, width=0.6)
    axes[0].axhline(0, color="#555555", linewidth=0.8)
    axes[0].set_title("Net return points", fontweight="bold", color="#17365D")
    axes[0].set_ylim(-17, 17)
    for i, v in enumerate(net): axes[0].text(i, v + (0.8 if v>=0 else -1.5), f"{v:+.2f}", ha="center", fontsize=9)
    axes[1].bar(labels, pf, color=colors, width=0.6)
    axes[1].axhline(1.0, color="#9B1C1C", linestyle="--", linewidth=1)
    axes[1].set_title("Profit factor", fontweight="bold", color="#17365D")
    axes[1].set_ylim(0, 2.05)
    for i, v in enumerate(pf): axes[1].text(i, v+0.05, f"{v:.3f}", ha="center", fontsize=9)
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.18)
    fig.suptitle("09:35 and 09:45 SHORT: candidates passing other setup filters", x=0.06, ha="left", fontweight="bold", color="#17365D")
    fig.tight_layout(rect=[0, 0, 1, 0.92])
    path3 = QA / "high_leg_pass_fail.png"
    fig.savefig(path3, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path1, path2, path3


def build():
    chart1, chart2, chart3 = charts()
    doc = configure_document()

    # First-page memo masthead.
    add_para(doc, "TECHNICAL VALIDATION MEMO", size=9.5, bold=True, color=BLUE, after=8)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("Is the OI Gate Useless?"), size=24, bold=True, color=NAVY)
    add_para(doc, "Independent review of Claude's V6 claim and its implications for V10 and V12", size=13, color=MUTED, after=16)
    meta = [
        ("Prepared for", "Strategy owner / backtest reviewer"),
        ("Prepared on", "2 September 2026"),
        ("Primary scope", "V6 selected trades, V6 signal cache, code path, and V10-V12 lineage"),
        ("Decision needed", "How to treat OI until a point-in-time rolling-contract test exists"),
    ]
    for label, value in meta:
        add_rich_para(doc, [
            (f"{label}: ", {"size": 10.5, "bold": True, "color": NAVY}),
            (value, {"size": 10.5, "color": INK}),
        ], after=3)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    add_callout(
        doc,
        "VERDICT",
        "Claude correctly identified a serious contract-age confound and correctly reproduced the winner/loser statistics. However, the statement that the OI gate does no discriminative work is not supported by those selected-trade tests and is partly contradicted by a direct per-leg threshold ablation. The defensible conclusion is: the present OI implementation is not cleanly validated as an economic signal, but it is not empirically inert in this replay.",
        fill=PALE_BLUE,
        color=NAVY,
    )

    add_heading(doc, "Executive conclusion", 1)
    add_bullet(doc, "The numerical claims from the 233-row V6 selected-trade file are accurate: 232 fills, 117 winners, 115 losers; winner median OI change 0.7895% versus loser median 0.9777%; Mann-Whitney p=0.7677; Welch p=0.2848; Spearman rho=-0.0227 (p=0.7309).")
    add_bullet(doc, "The data-quality diagnosis is real: every V6 selected trade uses a 26AUG future even though the sample starts on 27 May. Median previous OI rises about 682x from May to August, while median five-minute OI percentage change falls about 16.9x. OI change is strongly related to calendar position (Spearman rho=-0.406 versus date, p about 1.1e-10).")
    add_bullet(doc, "But winner-versus-loser correlation among already selected trades cannot test whether a threshold gate helped. Selection truncates OI values, mixes ten setups, and removes the rejected observations needed for a counterfactual.")
    add_bullet(doc, "When only the per-leg OI thresholds are relaxed to the existing 0.05% base floor, the replay changes materially: 233 to 272 orders, net return points 98.57 to 75.30, and PF 2.109 to 1.654. Therefore, the claim that the OI thresholds do no selection work, or that results should barely move, is false for this executable replay.")
    add_bullet(doc, "This does not prove OI has genuine forward edge. Thresholds were optimized in-sample, the base gate remains active in the ablation, and the OI scale is entangled with contract maturity. The right status is 'quarantined and unproven,' not 'useless' and not 'validated.'")

    add_heading(doc, "Claim-by-claim scorecard", 1)
    add_table(
        doc,
        ["Claude claim", "Assessment", "Independent finding"],
        [
            ("Selected OI does not separate winners and losers", "Accurate but narrow", "All quoted V6 tests reproduce. They describe the selected range only."),
            ("The metric is a calendar proxy", "Strongly supported", "Static 26AUG mapping plus large maturity drift; rho=-0.406 versus date."),
            ("The metric is structurally broken", "Overstated", "Percentage change is legitimate; this implementation/sample makes a fixed cutoff non-stationary."),
            ("Two short legs stop firing in August", "Accurate", "Zero selected 09:35/09:45 SHORT trades in August; all 15 other-filter candidates fail their high OI cutoffs."),
            ("The high thresholds did no useful work", "Contradicted in-sample", "Pass pool PF 1.788 versus fail pool PF 0.429 for the two high-threshold legs."),
            ("Removing OI should barely change results", "Not supported", "Per-leg threshold ablation loses 23.28 net points and 0.455 PF."),
            ("V10/V12 prove OI is useless", "Not established", "They inherit the gate and static core history, so their selected-trade null cannot identify gate value."),
        ],
        [2500, 1700, 5160],
        size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "1. What was tested", 1)
    add_heading(doc, "1.1 Evidence reviewed", 2)
    add_para(doc, "The review used the exact current local artifacts and code paths, not a verbal reconstruction:")
    for text in (
        "V6 trade audit: C:\\TradingData\\eqidv2\\fno_oi\\strategy_research\\ema_confirm_0925_0930_0935_0940_0945_v6_best_net_trades.csv (233 selections, 232 fills, 27 May to 21 August 2026).",
        "V6 pre-setup signal cache: ...\\_signal_cache_equity_1m_aggregated_5m_futures_oi_v4\\signals.parquet (12,436 candidates after the common base gate).",
        "V6 forward paths: ...\\_signal_cache_equity_1m_aggregated_5m_futures_oi_v4\\paths.npz.",
        "V11 and V12 locked reference-scenario closed trades and source_segments.json from their latest attested run directories.",
        "Implementation files: fno_oi_hybrid_data.py, fno_oi_ema_confirm_sweep.py, fno_v5_hybrid_backtest.py, fno_oi_ema_confirm_0925_0930_0935_0940_0945_v6.py, fno_v8_windowed_1m_entry_backtest.py, and the V10-V12 locked runners/registries.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "1.2 Three different questions that must not be mixed", 2)
    add_number(doc, "Does the continuous OI value correlate with outcome among trades already selected? Claude tested this. The answer is no in V6, V11, and V12.")
    add_number(doc, "Does the threshold reject worse candidates before ranking and entry? The selected-trade tests cannot answer this. The partial V6 ablation and pass/fail pool address it and show material in-sample separation.")
    add_number(doc, "Would OI retain incremental value with a point-in-time rolling near-month series on unseen data? No existing result in the reviewed lineage answers this cleanly.")

    add_callout(doc, "KEY LOGIC",
                "A gate can be useful even when values above the gate do not rank winners. Example: if everything below 1.0 is poor and everything above 1.0 is similar, the gate matters while the within-selected correlation is near zero. This is range restriction, not a contradiction.",
                fill=PALE_GOLD, color=GOLD)

    add_heading(doc, "2. What the code actually does", 1)
    add_heading(doc, "2.1 OI calculation and common base gate", 2)
    add_para(doc, "The hybrid join computes the five-minute feature as:")
    add_callout(doc, "FORMULA", "oi_change_pct = (current OI / previous five-minute OI - 1) x 100", fill=LIGHT, color=NAVY)
    add_para(doc, "The candidate builder then requires both current OI > previous OI and oi_change_pct >= 0.05%, in addition to the broad EMA, price-move, and volume gates. Consequently, the 12,436-row V6 signal cache has already excluded flat and falling OI. It is not a pre-OI universe.")

    add_heading(doc, "2.2 V6 setup-level thresholds", 2)
    add_para(doc, "V6 applies a second OI threshold inside each setup. Eight legs use 0.10% or 0.25%. Two short legs are much stricter:")
    add_table(
        doc,
        ["Signal / entry", "Side", "OI threshold", "August selected trades"],
        [
            ("09:35 / 09:36", "SHORT", ">= 1.00%", "0"),
            ("09:45 / 09:46", "SHORT", ">= 0.75%", "0"),
        ],
        [2600, 1500, 2500, 2760],
        numeric_cols=(2, 3),
        size=9.5,
    )
    add_para(doc, "For FILTERED setups, V6 first applies all thresholds, then ranks qualifying rows by the setup picker and takes a daily cap. Changing the OI cutoff can therefore alter both the admitted names and the name selected by a picker.")

    add_heading(doc, "2.3 Propagation into later generations", 2)
    add_table(
        doc,
        ["Generation", "OI behavior", "Contract-history issue", "Implication"],
        [
            ("V6", "Base >0 and >=0.05%; per-leg cutoffs", "26AUG static universe across May-Aug", "Both base and leg gates are confounded by maturity."),
            ("V8", "Independent builder repeats same base gate and setup cutoff", "Static source snapshot by run", "Reimplementation preserves the same selection seam."),
            ("V10 / V11", "Consume V8-style gated candidate caches", "Static universes by segment; core 59 sessions use 26AUG", "Later runtime changes do not recover rejected OI rows."),
            ("V12", "Starts from 1,241 already gated all-input candidates", "Registry explicitly flags non-point-in-time core history", "Selected-trade feature tests remain conditional on admission."),
        ],
        [1500, 2420, 2690, 2750],
        size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "3. Reproduction of Claude's selected-trade statistics", 1)
    add_heading(doc, "3.1 V6 winners versus losers", 2)
    add_table(
        doc,
        ["Measure", "Winners", "Losers", "Test / interpretation"],
        [
            ("Filled trades", "117", "115", "232 total; one selection did not fill"),
            ("Median oi_change_pct", "0.7895%", "0.9777%", "Loser median is slightly higher"),
            ("Mann-Whitney U", "", "", "U=6,576; p=0.7677"),
            ("Welch t-test", "", "", "t=1.074; p=0.2848"),
            ("Spearman versus net return", "", "", "rho=-0.0227; p=0.7309"),
            ("Rank-biserial effect", "", "", "about -0.023; negligible"),
        ],
        [2500, 1550, 1550, 3760],
        numeric_cols=(1, 2),
        size=9,
    )
    add_para(doc, "This part of Claude's analysis is correct. It says that, inside the admitted and daily-selected V6 trades, a larger OI percentage increase is not a monotonic predictor of a better trade return.")
    add_para(doc, "It does not say that the cutoff did nothing. The observations below the setup cutoff are mostly absent from this table by construction, and all observations at or below the common base gate are absent from the signal cache itself.")

    add_heading(doc, "3.2 The same narrow null appears in V11 and V12", 2)
    add_table(
        doc,
        ["Strategy", "Wins / losses", "Winner median", "Loser median", "Mann-Whitney p", "Spearman rho (p)"],
        [
            ("V6", "117 / 115", "0.7895%", "0.9777%", "0.7677", "-0.0227 (0.7309)"),
            ("V11", "123 / 114", "0.9120%", "0.9060%", "0.5863", "+0.0191 (0.7703)"),
            ("V12", "120 / 109", "0.9016%", "0.8827%", "0.6793", "+0.0288 (0.6649)"),
        ],
        [1350, 1550, 1550, 1550, 1600, 1760],
        numeric_cols=(1, 2, 3, 4, 5),
        size=8.5,
    )
    add_para(doc, "This consistency strengthens the conclusion that raw OI magnitude above the gate is not a useful continuous ranker among filled trades. It still does not identify the gate's incremental effect because V11/V12 never reconstruct the candidates excluded by the common OI gate.")

    add_heading(doc, "4. Contract-age distortion", 1)
    add_heading(doc, "4.1 Static contract mapping is confirmed", 2)
    add_para(doc, "All 233 V6 selections map to a 26AUG future expiring 25 August 2026, while the selected sample starts 27 May. The 11 August dated universe contains 213 August contracts. Thus the same expiry cohort is read backward from far-month to front-month status.")

    add_table(
        doc,
        ["Month", "Selected rows", "Median previous OI", "Median OI change"],
        [
            ("May 2026", "4", "16,988", "7.2751%"),
            ("June 2026", "67", "135,750", "1.6949%"),
            ("July 2026", "117", "1,944,800", "0.8827%"),
            ("August 2026", "45", "11,583,000", "0.4304%"),
        ],
        [1900, 1500, 2800, 3160],
        numeric_cols=(1, 2, 3),
        size=9.5,
    )
    add_picture(
        doc,
        chart1,
        "Figure 1. Monthly medians from the 233 V6 selections. The OI base rises about 682x; median percentage change falls about 16.9x.",
        "Bar and line chart showing median previous OI increasing sharply from May through August while median five-minute OI percentage change declines.",
    )

    add_heading(doc, "4.2 Statistical relationship with time", 2)
    add_bullet(doc, "oi_change_pct versus calendar date: Spearman rho=-0.4065, p=1.10e-10.")
    add_bullet(doc, "previous OI versus calendar date: Spearman rho=+0.6888, p=4.10e-34.")
    add_bullet(doc, "OI distributions differ by month: Kruskal-Wallis H=48.92, p=1.35e-10.")
    add_bullet(doc, "After rank-residualizing OI and return for date and setup, the partial rank correlation remains small: about -0.061 (p=0.358).")
    add_para(doc, "These are strong indicators of non-stationarity. They do not prove that every movement in OI is caused by contract age; symbol mix, day regime, and selection also matter. But they are more than sufficient to reject the idea that one absolute percentage cutoff has a stable meaning across this history.")

    add_heading(doc, "4.3 Important unit correction", 2)
    add_callout(doc, "CORRECTION",
                "The examples are quantities, not counts of futures contracts. MFSL previous OI=400 equals one 400-unit lot; current OI=1,200 equals three lots. ONGC previous OI=9,000 equals four 2,250-unit lots; current OI=15,750 equals seven lots. The small-base/discreteness concern remains valid, but calling these '400 contracts' or '9,000 contracts' overstates market depth.",
                fill=PALE_RED, color=RED)
    add_para(doc, "Percentage normalization is not inherently wrong: the same absolute addition should often matter less in a deeper market. The problem is that the denominator, minimum lot increment, and contract participation change systematically with maturity, while the cutoff is fixed and was optimized on this same life cycle.")

    add_page_break(doc)
    add_heading(doc, "5. Direct V6 threshold ablation", 1)
    add_heading(doc, "5.1 Design", 2)
    add_para(doc, "I reran the existing V6 setup-selection and bracket simulator from the cached 12,436-signal superset. The only change was to set every setup-level OI threshold to the cache's existing 0.05% floor. The common base requirements (OI rising and >=0.05%) remained active, so this is a per-leg threshold ablation, not a complete no-OI test.")
    add_para(doc, "All other setup thresholds, pickers, daily caps, stops, targets, forward paths, and five-basis-point round-trip cost were unchanged.")

    add_table(
        doc,
        ["Replay", "Orders / fills", "Wins / losses", "Net points", "PF", "Change versus frozen"],
        [
            ("Frozen V6 thresholds", "233 / 232", "117 / 115", "98.574", "2.109", "Baseline"),
            ("All leg OI thresholds = 0.05%", "272 / 270", "124 / 146", "75.297", "1.654", "-23.277 net; -0.455 PF"),
        ],
        [2250, 1550, 1450, 1350, 1100, 1660],
        numeric_cols=(1, 2, 3, 4),
        size=9,
    )
    add_picture(
        doc,
        chart2,
        "Figure 2. Lowering every setup OI cutoff to 0.05% materially changes the V6 replay. The common base OI gate is still active.",
        "Two-panel bar chart comparing frozen V6 thresholds with setup thresholds lowered to 0.05 percent; net return and profit factor are both lower under the relaxed thresholds.",
    )

    add_heading(doc, "5.2 Period breakdown", 2)
    add_table(
        doc,
        ["Period", "Frozen net / PF", "0.05% net / PF", "Interpretation"],
        [
            ("Before 17 Jul optimizer split", "64.895 / 2.295", "60.979 / 2.080", "Small deterioration"),
            ("17 Jul onward", "33.679 / 1.868", "14.318 / 1.244", "Large deterioration"),
            ("Through 11 Aug selection history", "101.153 / 2.297", "78.997 / 1.801", "Large in-sample advantage"),
            ("Strictly after 11 Aug", "-2.579 / 0.763", "-3.700 / 0.776", "Both lose; no forward validation"),
            ("August, all available", "7.945 / 1.406", "-2.149 / 0.937", "Relaxed book adds weak trades"),
        ],
        [2150, 1850, 1850, 3510],
        numeric_cols=(1, 2),
        size=8.8,
    )
    add_callout(doc, "INTERPRETATION",
                "This ablation refutes 'the thresholds do nothing.' It does not establish a causal OI edge. The apparent benefit can be a fitted time/regime filter created by contract maturity. The strict post-selection sample is only 24 versus 35 trades and both variants lose.",
                fill=PALE_GOLD, color=GOLD)

    add_heading(doc, "6. The two high-threshold short legs", 1)
    add_heading(doc, "6.1 Did the thresholds switch them off in August?", 2)
    add_para(doc, "Yes. Within the already base-gated signal cache, candidates that passed every other setup filter were:")
    add_table(
        doc,
        ["Leg", "Month", "Other-filter pool", "Passed OI cutoff", "Failed OI cutoff", "Median OI change"],
        [
            ("09:35 SHORT >=1.00%", "Jun", "15", "9", "6", "1.4706%"),
            ("", "Jul", "18", "7", "11", "0.8562%"),
            ("", "Aug", "7", "0", "7", "0.2690%"),
            ("09:45 SHORT >=0.75%", "May", "1", "1", "0", "7.1429%"),
            ("", "Jun", "8", "7", "1", "2.8266%"),
            ("", "Jul", "19", "13", "6", "1.0638%"),
            ("", "Aug", "8", "0", "8", "0.1782%"),
        ],
        [2350, 850, 1550, 1500, 1500, 1610],
        numeric_cols=(1, 2, 3, 4, 5),
        size=8.3,
    )
    add_para(doc, "All 15 August candidates fail the setup-level OI cutoffs. When those two cutoffs are relaxed to 0.05%, daily ranking selects 12 August trades; they produce -3.067 net points, PF 0.583, and a 5/7 win/loss count. In this historical sample, switching the legs off avoided losses.")

    add_heading(doc, "6.2 Pass versus fail outcomes before the daily cap", 2)
    add_table(
        doc,
        ["Group", "Candidates / fills", "Win rate", "Net points", "PF"],
        [
            ("Passed high OI threshold", "37 / 37", "48.6%", "+12.864", "1.788"),
            ("Failed high OI threshold", "39 / 37", "35.1%", "-13.613", "0.429"),
        ],
        [2800, 1900, 1500, 1600, 1560],
        numeric_cols=(1, 2, 3, 4),
        size=9.5,
    )
    add_picture(
        doc,
        chart3,
        "Figure 3. For the two high-threshold short legs, candidates passing all other filters perform much better above the OI cutoff in this in-sample, base-gated pool.",
        "Two-panel bar chart showing positive net return and profit factor above one for high-leg candidates that pass the OI cutoff, versus negative net and profit factor below one for those that fail.",
    )
    add_para(doc, "For 09:35 SHORT alone, pass-versus-fail returns produce an unadjusted Mann-Whitney p=0.0375. That result is exploratory: it is one of ten legs, thresholds were optimized, and no multiple-testing correction or independent holdout makes it confirmatory. For 09:45 SHORT, both pass and fail pools lose, although the fail pool is worse.")

    add_page_break(doc)
    add_heading(doc, "7. Why both sets of findings can be true", 1)
    add_heading(doc, "7.1 Selected-trade null versus gate effect", 2)
    add_para(doc, "Claude measured whether larger OI values among selected fills rank outcomes. The ablation measures whether admitting lower-OI candidates changes the book. These are different estimands. A threshold can remove a bad tail while leaving no relationship inside the retained range.")
    add_para(doc, "There is an additional portfolio interaction: the daily cap and picker mean a newly admitted low-OI candidate can displace a previously selected candidate. In the relaxed replay, 58 selections are new and 19 original selections are displaced. The 58 new selections produce -14.789 net points and PF 0.547; the 19 displaced selections had produced +8.488 net points and PF 2.337.")

    add_heading(doc, "7.2 Economic signal versus accidental time filter", 2)
    add_para(doc, "The ablation shows that the numeric threshold affected historical profitability. It cannot tell us why. Two mechanisms can create the same result:")
    add_bullet(doc, "Economic mechanism: unusually strong positive OI expansion genuinely identifies fresh participation that supports continuation.")
    add_bullet(doc, "Artifact mechanism: the threshold selects earlier contract-life periods or market regimes that happened to perform better, while rejecting later periods because the denominator is larger.")
    add_para(doc, "Because the core history uses one later contract cohort backward, both mechanisms are mixed. Therefore, saying 'OI increased profitability' is premature even though the threshold improved the in-sample replay.")

    add_heading(doc, "8. Corrected answer in simple language", 1)
    add_callout(doc, "PLAIN ENGLISH",
                "OI is not proven useless. The current test is like using one ruler that changes scale every month. Claude correctly noticed the scale problem. But when the OI cutoffs are loosened, the strategy actually gets worse in this replay, so the cutoffs were doing something. We still do not know whether they found real fresh-money behavior or merely filtered dates because the August contract matured. Treat OI as unverified, not worthless.",
                fill=PALE_GREEN, color=GREEN)
    add_table(
        doc,
        ["Question", "Best current answer"],
        [
            ("Is OI valuable in markets?", "Potentially yes; the concept is economically plausible."),
            ("Is raw five-minute percentage OI cleanly measured here?", "No. A static contract cohort creates strong maturity non-stationarity."),
            ("Did the current OI thresholds affect V6 selection/profit?", "Yes, materially, in this historical replay."),
            ("Does that prove causal or forward edge?", "No. Optimization and contract-age confounding remain."),
            ("Do V10/V12 solve the problem?", "No. They inherit already OI-gated candidate caches and the static core history."),
            ("Should OI be removed immediately?", "Not on the selected-trade p-values alone. Run a clean rolling-contract ablation first."),
        ],
        [3200, 6160],
        size=9.2,
    )

    add_heading(doc, "9. Recommended validation program", 1)
    add_heading(doc, "9.1 Minimum decisive experiment: predeclared 2 x 2 ablation", 2)
    add_para(doc, "Build the candidate stream before any OI condition and run four frozen variants on identical price/volume data, ranking, caps, execution rules, and costs:")
    add_table(
        doc,
        ["Variant", "Base OI direction/floor", "Per-leg OI cutoffs", "Question answered"],
        [
            ("A", "ON", "ON", "Current lineage control"),
            ("B", "ON", "OFF", "Incremental value of optimized leg thresholds"),
            ("C", "OFF", "ON where defined", "Whether direction/base floor matters separately"),
            ("D", "OFF", "OFF", "Complete no-OI counterfactual"),
        ],
        [1100, 2600, 2400, 3260],
        size=9,
    )
    add_para(doc, "The primary comparison should be A versus D. B and C diagnose which layer contributes. The current review executes only A versus B with B floored at 0.05%, because the stored V6 cache has already removed base-gate failures.")

    add_heading(doc, "9.2 Data construction required", 2)
    add_number(doc, "Use a point-in-time contract mapping for every session. Define and freeze the rollover rule before outcomes are examined (for example, actual front month with an explicit rollover date).")
    add_number(doc, "Retain the full pre-OI candidate stream, including negative, zero, missing, and small positive OI changes, with rejection reasons.")
    add_number(doc, "Express raw OI in both units and lots: delta_oi_lots = (oi - prev_oi) / dated_lot_size. Preserve the raw denominator and days-to-expiry.")
    add_number(doc, "Test stable normalizations: within-symbol rolling z-score of delta OI; delta OI divided by rolling median OI; and cross-sectional percentile within a date/slot/DTE bucket. Do not tune all of them on the same holdout.")
    add_number(doc, "Add persistence only with causal observations (for example, two completed five-minute bars). V12 already records that the current cache lacks the required sidecar for this validity test.")

    add_heading(doc, "9.3 Evaluation and promotion gates", 2)
    add_bullet(doc, "Freeze the strategy and OI definition before the forward period. Report train, validation, and genuinely untouched forward results separately.")
    add_bullet(doc, "Measure affected decisions, not only filled trades: candidate admissions, re-rankings, displacements, fills, side/leg/month effects, and portfolio overlap.")
    add_bullet(doc, "Use paired day-level bootstrap confidence intervals for net and PF differences; use false-discovery control for per-feature/per-leg tests.")
    add_bullet(doc, "Require stability across contract months and days-to-expiry buckets. A cutoff that works only when the contract is far month or only in one calendar month fails.")
    add_bullet(doc, "A practical minimum: at least 30 affected decisions in forward data, non-negative delta under all declared cost scenarios, no material side collapse, and no dependence on one best day.")

    add_heading(doc, "10. Decision recommendation", 1)
    add_callout(doc, "RECOMMENDATION",
                "Do not advertise OI as a validated profitability driver in V6, V10, or V12. Also do not delete it solely because selected winners and losers have similar OI values. Freeze the current OI book as a control, label it maturity-confounded, and prioritize a rolling-contract A/B/C/D ablation. Until it passes forward validation, treat OI as a quarantined selection feature rather than strategy thesis.",
                fill=PALE_BLUE, color=NAVY)
    add_para(doc, "Operationally, if a production decision must be made before that experiment, the safer interpretation is that the strategy's demonstrated ingredients are its price/EMA/volume/confirmation/execution stack; OI is an unresolved overlay. Any expansion beyond the F&O universe should be tested as a new strategy variant rather than inferred from p=0.77.")

    add_page_break(doc)
    add_heading(doc, "Appendix A. Reproducibility details", 1)
    add_heading(doc, "A.1 Exact V6 ablation parameters", 2)
    add_table(
        doc,
        ["Element", "Frozen control", "Ablated replay"],
        [
            ("Signal cache", "12,436 base-gated rows", "Same"),
            ("Common OI gate", "OI > previous; change >=0.05%", "Same"),
            ("Setup OI thresholds", "0.10%-1.00% by leg", "0.05% for every leg"),
            ("Price/volume/body/wick", "Frozen V6", "Same"),
            ("Picker and daily caps", "Frozen V6", "Same"),
            ("Stops and targets", "Frozen V6", "Same"),
            ("Round-trip cost", "5 bps", "Same"),
            ("Forward paths", "Cached one-minute paths", "Same"),
        ],
        [2400, 3480, 3480],
        size=9,
    )

    add_heading(doc, "A.2 Key code locations", 2)
    add_table(
        doc,
        ["File", "Relevant responsibility"],
        [
            ("fno_oi_hybrid_data.py:410-434", "Computes previous OI and percentage change, then joins only OI fields to cash features."),
            ("fno_oi_ema_confirm_sweep.py:52, 233-236", "Defines 0.05% loose/base floor and requires rising OI in V6 signal construction."),
            ("fno_v5_hybrid_backtest.py:29-47", "Applies per-setup OI threshold with price, volume, body, wick, and liquidity gates."),
            ("..._v6.py:135-145", "Freezes the ten V6 setup thresholds, including 1.00% and 0.75% short legs."),
            ("fno_v8_windowed_1m_entry_backtest.py:171-173, 2121-2154", "Repeats the 0.05% base OI gate and setup-specific OI threshold in the independent builder."),
            ("fno_v12_variant_registry.py:507-514", "Flags the non-point-in-time static futures universe as a blocked validity test."),
        ],
        [4100, 5260],
        size=8.7,
    )

    add_heading(doc, "A.3 Scope limits", 2)
    add_bullet(doc, "The full no-OI base-gate counterfactual was not run because the stored V6 signal cache starts after the common OI gate. Rebuilding it is the next decisive test.")
    add_bullet(doc, "The V6 current-source replay (233 orders; net 98.57; PF 2.109) is not the historical 11 August headline artifact (about +144%; PF about 2.80). The report does not blend those result sets.")
    add_bullet(doc, "Candidate return comparisons reuse optimized thresholds and historical paths. They describe this sample; they are not independent evidence of live profitability.")
    add_bullet(doc, "No claim here is investment advice. This is a validation of code, data construction, and statistical interpretation.")

    add_heading(doc, "Appendix B. Final answer", 1)
    add_para(doc, "Validated conclusion, in one sentence:", bold=True, color=NAVY)
    add_callout(doc, "BOTTOM LINE",
                "Claude found a real maturity bias and a real lack of within-selected OI correlation, but overreached by calling the gate useless: the per-leg OI cutoffs materially improved this in-sample V6 replay, while the current static-contract design prevents us from knowing whether that improvement is genuine OI information or an accidental calendar filter.",
                fill=PALE_GREEN, color=GREEN)

    # Core properties and save.
    doc.core_properties.title = "OI Gate Validation: V6, V10 and V12"
    doc.core_properties.subject = "Independent validation of OI gate usefulness and maturity confounding"
    doc.core_properties.author = "Independent technical review"
    doc.core_properties.keywords = "open interest, V6, V10, V12, ablation, backtest validation"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
